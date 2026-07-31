"""Build the submission-ready CareTrust Track 2 Phase 1 DOCX.

The source of truth is the Markdown narrative. The document uses the
``grant_proposal`` design preset with one named compactness override required by
the challenge: Arial 11 pt, single-spaced narrative, and 4 pt paragraph space.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submission" / "CareTrust_Track2_Phase1_Application_Draft.md"
OUTPUT = ROOT / "submission" / "CareTrust_Track2_Phase1_Application_Submission.docx"

FONT = "Arial"
INK = "17211B"
BLUE = "1F4D78"
BLUE_LIGHT = "E8EEF5"
MUTED = "59645E"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float = 11, bold: bool | None = None,
                 italic: bool | None = None, color: str = INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run_node = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run_node.extend([run_props, text_node])
    hyperlink.append(run_node)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, *, size: float = 11, base_bold: bool = False,
               base_italic: bool = False, color: str = INK) -> None:
    for part in filter(None, INLINE_RE.split(text)):
        link_match = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", part)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=base_italic, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, bold=base_bold, italic=base_italic, color=color)
            run.font.name = "Courier New"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=base_bold, italic=base_italic, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, 14, 7),
        "Heading 2": (13, 10, 5),
        "Heading 3": (12, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0


def configure_section(section, *, narrative: bool) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    if narrative:
        section.header.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(header, "CareTrust | Track 2 Phase 1", size=9, base_bold=True, color=MUTED)
        header.paragraph_format.space_after = Pt(0)
        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_inline(footer, "Project Narrative | ", size=9, color=MUTED)
        add_page_field(footer)
        pg_num = OxmlElement("w:pgNumType")
        pg_num.set(qn("w:start"), "1")
        section._sectPr.append(pg_num)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = len(rows[0])
    if columns == 2:
        widths = [2600, TABLE_WIDTH_DXA - 2600]
    elif columns == 3:
        widths = [1800, 3000, 4560]
    else:
        base = TABLE_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(
                paragraph,
                value,
                size=9,
                base_bold=row_index == 0,
                color=INK,
            )
            if row_index == 0:
                shade_cell(cell, BLUE_LIGHT)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.line_spacing = 1.0


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        if not is_table_separator(lines[index]):
            rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], narrative=False)
    doc.core_properties.title = "CareTrust - Caregiver AI Challenge Track 2 Phase 1 Application"
    doc.core_properties.subject = "AI Tools for Extending the Caregiver Workforce"
    doc.core_properties.author = "Michael Makani Kai McDougall"
    doc.core_properties.keywords = "CareTrust, caregiver workforce, AI, interoperability, Track 2"

    in_cover = True
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line == "[PAGE BREAK]":
            if in_cover:
                narrative = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(narrative, narrative=True)
                in_cover = False
            else:
                doc.add_page_break()
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2), size={1: 16, 2: 13, 3: 12}[level], base_bold=True, color=BLUE)
            if in_cover and level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(10)
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline(paragraph, line[2:], base_italic=True, color=MUTED)
            index += 1
            continue
        paragraph = doc.add_paragraph()
        if in_cover and ":" in line and not line.startswith("http"):
            label, value = line.split(":", 1)
            add_inline(paragraph, f"{label}:", base_bold=True)
            add_inline(paragraph, value)
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            add_inline(paragraph, line)
        index += 1

    # Ensure the cover does not inherit narrative page furniture.
    first = doc.sections[0]
    first.header.paragraphs[0].clear()
    first.footer.paragraphs[0].clear()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
