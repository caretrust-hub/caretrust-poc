"""Create submission-safe Markdown and DOCX logs from frozen Smart40 evidence."""
from __future__ import annotations
import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
RUN=ROOT/'artifacts'/'validation'/'intent-compiler-bedrock-40'
OUT=ROOT/'submission'
MD=OUT/'Smart40-Data-Output-Logs.md'
DOCX=OUT/'Smart40-Data-Output-Logs.docx'

DISCLOSURE="Protocol defect disclosure: the frozen runner required citation span_ids but sent only utterance text to the model; valid span IDs were not supplied. Candidate validation 0/39 therefore reflects this evaluator/protocol mismatch as well as any model error. The frozen sequence and raw JSONL were not changed."


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 103, 109)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])

def load():
    frozen=json.loads((RUN/'frozen-config.json').read_text())
    rows=[json.loads(line) for line in (RUN/'results.jsonl').read_text().splitlines() if line]
    if len(rows)!=40 or [r['ordinal'] for r in rows]!=list(range(1,41)): raise ValueError('Smart40 must contain exact ordered 40 records')
    if frozen['freeze_sha256']!=json.loads((RUN/'summary.json').read_text())['frozen_config_sha256']: raise ValueError('frozen hash mismatch')
    return frozen,rows

def markdown(frozen,rows):
    parts=['# CareTrust Smart40 Data Output Logs','\n## Cover and method note',f'Frozen configuration SHA-256: `{frozen["freeze_sha256"]}`.','This is a judge-readable rendering of the retained frozen Smart40 evaluation. It is not a replacement for raw JSONL, which remains separately retained and is not claimed as a submission attachment.','All 40 ordered cycles below are reproduced from the retained JSONL. Errors are retained, not omitted.',f'\n**{DISCLOSURE}**','\n## Ordered cycles']
    for row in rows: parts += [f'\n### Cycle {row["ordinal"]}: {row["case_id"]} ({row["group"]})','```json',json.dumps(row,indent=2,sort_keys=True),'```']
    return '\n'.join(parts)+'\n'

def build():
    frozen,rows=load(); OUT.mkdir(exist_ok=True); text=markdown(frozen,rows); MD.write_text(text,encoding='utf-8')
    doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.65); sec.left_margin=sec.right_margin=Inches(.7)
    sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)
    style=doc.styles['Normal']; style.font.name='Arial'; style.font.size=Pt(10.5)
    style.paragraph_format.space_after=Pt(4); style.paragraph_format.line_spacing=1.0
    for heading_name, size in (("Title", 22), ("Heading 1", 15)):
        heading = doc.styles[heading_name]
        heading.font.name = "Arial"
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(23, 58, 75)
        heading.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header_run = header.add_run("CARETRUST  |  SMART40 FROZEN DATA OUTPUT LOGS")
    header_run.font.name = "Arial"; header_run.font.size = Pt(8.5); header_run.bold = True
    header_run.font.color.rgb = RGBColor(90, 103, 109)
    add_page_number(sec.footer.paragraphs[0])
    doc.add_heading('CareTrust Smart40 Data Output Logs',0); doc.add_paragraph('Frozen consecutive Bedrock evaluation - judge submission log')
    doc.add_heading('Cover and method note',1); doc.add_paragraph(f'Frozen configuration SHA-256: {frozen["freeze_sha256"]}')
    results_sha256=hashlib.sha256((RUN/'results.jsonl').read_bytes()).hexdigest()
    doc.add_paragraph(f'Retained results.jsonl SHA-256: {results_sha256}')
    doc.add_paragraph('This Word log renders all retained cycles in order. The raw JSONL is separately retained and is not represented as a submission attachment.')
    doc.add_paragraph('Safety/error retention: every cycle is shown, including the retained safety error at cycle 34; no retry, reordering, or omission occurred.')
    doc.add_paragraph('Run summary: 40/40 ordered records retained; 39 candidate responses and one deterministic safety error; estimated model cost $0.00391725; all downstream outputs remained draft-only.')
    doc.add_paragraph('Observed fields: action 39/39; resource 39/39; purpose 39/39; audience 38/39; expiry 33/39. These are field observations, not autonomous-authorization claims.')
    disclosure=doc.add_paragraph(DISCLOSURE); disclosure.runs[0].bold=True
    for row in rows:
        doc.add_page_break(); doc.add_heading(f'Cycle {row["ordinal"]}: {row["case_id"]} ({row["group"]})',1)
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
        r=p.add_run(json.dumps(row,indent=2,sort_keys=True)); r.font.name='Consolas'; r.font.size=Pt(10)
    doc.core_properties.title='CareTrust Smart40 Data Output Logs'
    doc.core_properties.subject='ACL Caregiver AI Challenge Track 2 data output log attachment'
    doc.core_properties.author='Michael Makani Kai McDougall'
    doc.core_properties.keywords='CareTrust, Smart40, Bedrock, data output logs, responsible AI'
    doc.save(DOCX); return MD,DOCX
if __name__=='__main__':
    for path in build(): print(path.relative_to(ROOT))
